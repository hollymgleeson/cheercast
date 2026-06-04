from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Page margins ---
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = int(8.5 * 914400)
section.page_height = int(11  * 914400)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# --- Default style ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def body(content="", bold_prefix=None, space_before=0, space_after=6, text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    actual_text = text if text else content
    if bold_prefix:
        run = p.add_run(bold_prefix + "  ")
        run.bold = True
        run.font.name = 'Calibri'
    run = p.add_run(actual_text)
    run.font.name = 'Calibri'
    return p

def blank_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

def sig_line(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{'_' * 45}   {label}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

def sig_field(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(label)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============================================================
# TITLE BLOCK
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
r = title.add_run("CHEERCAST PILOT PROGRAM AGREEMENT")
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
r2 = sub.add_run("Confidentiality, Intellectual Property & Beta Testing Agreement")
r2.font.size = Pt(10)
r2.font.name = 'Calibri'
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

blank_line()

# Effective date / parties block
parties = doc.add_paragraph()
parties.paragraph_format.space_after = Pt(12)
parties.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = parties.add_run(
    "This Pilot Program Agreement (\"Agreement\") is entered into as of ________________, 2025 "
    "(\"Effective Date\") by and between:\n\n"
    "Gleeson Consulting, LLC, a limited liability company, located at 1215 Cedar Grove Road, Media, PA 19063 "
    "(\"Company\"), and\n\n"
    "Kedron Cheerleading Association, located at 3932 Miller Road, Newtown Square, PA 19073 (\"Gym\").\n\n"
    "Company and Gym are each referred to herein individually as a \"Party\" and collectively as the \"Parties.\""
)
r.font.name = 'Calibri'
r.font.size = Pt(11)

# ============================================================
# RECITALS
# ============================================================
heading("BACKGROUND", level=1)
body(
    "Company has developed a proprietary software platform known as CheerCast (the \"Platform\"), "
    "designed to assist cheerleading gyms with athlete management, team placement, evaluations, and "
    "related operational functions. Company wishes to conduct a limited, free-of-charge pilot program "
    "with Gym to test and refine the Platform. Gym wishes to participate in the pilot on the terms set "
    "forth in this Agreement."
)

# ============================================================
# SECTIONS
# ============================================================

heading("1.  PILOT PROGRAM", level=1)
body("", bold_prefix="1.1  Scope.",
     text="Company grants Gym a limited, non-exclusive, non-transferable, revocable license to access "
          "and use the Platform solely for internal evaluation purposes during the Pilot Period defined below. "
          "No rights are granted beyond those expressly stated in this Agreement.")
body("", bold_prefix="1.2  Pilot Period.",
     text="The pilot program will begin on the Effective Date and continue for ______ months, unless "
          "earlier terminated by either Party upon seven (7) days' written notice, or by Company immediately "
          "for cause.")
body("", bold_prefix="1.3  No Charge.",
     text="Gym's participation in the pilot is provided free of charge. Company reserves the right to "
          "introduce fees for continued use of the Platform after the Pilot Period, at its sole discretion.")
body("", bold_prefix="1.4  Beta Nature of Platform.",
     text="Gym acknowledges that the Platform is pre-release software provided on an \"as-is\" basis. "
          "It may contain errors, bugs, or incomplete features. Company does not warrant uninterrupted "
          "or error-free operation during the pilot.")

heading("2.  CONFIDENTIALITY", level=1)
body("", bold_prefix="2.1  Confidential Information.",
     text="\"Confidential Information\" means any non-public information disclosed by Company to Gym in "
          "connection with this Agreement, including but not limited to: the Platform's design, features, "
          "functionality, source code, algorithms, workflows, data structures, pricing strategy, business "
          "plans, user interface, and any other technical or commercial information. Confidential Information "
          "includes information disclosed orally, in writing, or by any other means.")
body("", bold_prefix="2.2  Obligations.",
     text="Gym agrees to: (a) hold all Confidential Information in strict confidence; "
          "(b) not disclose Confidential Information to any third party without the prior written consent "
          "of Company; (c) use Confidential Information solely to participate in the pilot program; and "
          "(d) limit access to Confidential Information to Gym's own personnel who have a need to know "
          "and who are bound by confidentiality obligations at least as protective as those in this Agreement.")
body("", bold_prefix="2.3  Duration.",
     text="Gym's confidentiality obligations survive the expiration or termination of this Agreement for "
          "a period of five (5) years.")
body("", bold_prefix="2.4  Exceptions.",
     text="The obligations in Section 2.2 do not apply to information that: (a) is or becomes publicly "
          "known through no breach by Gym; (b) was rightfully known to Gym before disclosure by Company; "
          "or (c) is required to be disclosed by law or court order, provided Gym gives Company prompt "
          "written notice and reasonably cooperates in seeking a protective order.")

heading("3.  INTELLECTUAL PROPERTY", level=1)
body("", bold_prefix="3.1  Ownership.",
     text="The Platform and all related intellectual property — including software, code, design, "
          "trademarks, trade secrets, and all improvements, modifications, or derivative works — are and "
          "shall remain the exclusive property of Company. This Agreement does not transfer any ownership "
          "interest in the Platform to Gym.")
body("", bold_prefix="3.2  No Reverse Engineering.",
     text="Gym agrees not to: (a) copy, reproduce, modify, or create derivative works based on the "
          "Platform; (b) reverse engineer, decompile, disassemble, or otherwise attempt to derive the "
          "source code or underlying structure of the Platform; or (c) remove, obscure, or alter any "
          "proprietary notices or labels on the Platform.")
body("", bold_prefix="3.3  Feedback.",
     text="If Gym provides suggestions, feedback, or ideas regarding the Platform (\"Feedback\"), Gym "
          "hereby assigns to Company all right, title, and interest in and to such Feedback. Company may "
          "use Feedback for any purpose without restriction or compensation to Gym.")
body("", bold_prefix="3.4  Gym Data.",
     text="Gym retains ownership of the data it inputs into the Platform (\"Gym Data\"). Company will not "
          "sell or disclose Gym Data to third parties without Gym's consent, except as required by law. "
          "Company may use anonymized, aggregated data derived from Gym Data to improve the Platform.")

heading("4.  PERMITTED USE & RESTRICTIONS", level=1)
body("", bold_prefix="4.1  Permitted Use.",
     text="Gym may use the Platform solely for its own internal gym operations during the Pilot Period.")
body("", bold_prefix="4.2  Restrictions.",
     text="Gym shall not: (a) sublicense, sell, resell, transfer, or otherwise make the Platform "
          "available to any third party; (b) use the Platform to develop a competing product or service; "
          "(c) use the Platform in any unlawful manner; or (d) share login credentials or access with "
          "anyone outside Gym's own organization.")

heading("5.  LIMITATION OF LIABILITY", level=1)
body(
    "TO THE MAXIMUM EXTENT PERMITTED BY LAW, COMPANY SHALL NOT BE LIABLE TO GYM FOR ANY INDIRECT, "
    "INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS "
    "AGREEMENT OR THE USE OF THE PLATFORM, EVEN IF COMPANY HAS BEEN ADVISED OF THE POSSIBILITY OF "
    "SUCH DAMAGES. COMPANY'S TOTAL LIABILITY TO GYM UNDER THIS AGREEMENT SHALL NOT EXCEED ZERO "
    "DOLLARS ($0), REFLECTING THE FREE-OF-CHARGE NATURE OF THIS PILOT."
)

heading("6.  TERM & TERMINATION", level=1)
body("", bold_prefix="6.1  Term.",
     text="This Agreement commences on the Effective Date and continues until the end of the Pilot "
          "Period, unless earlier terminated.")
body("", bold_prefix="6.2  Termination.",
     text="Either Party may terminate this Agreement at any time upon seven (7) days' written notice. "
          "Company may terminate immediately if Gym breaches any provision of this Agreement.")
body("", bold_prefix="6.3  Effect of Termination.",
     text="Upon termination, Gym's right to access and use the Platform immediately ceases. Gym shall "
          "promptly delete or return any Confidential Information in its possession. Sections 2, 3, 5, "
          "and 7 survive termination.")

heading("7.  GENERAL PROVISIONS", level=1)
body("", bold_prefix="7.1  Governing Law.",
     text="This Agreement is governed by the laws of the Commonwealth of Pennsylvania, without regard to "
          "conflict of law principles.")
body("", bold_prefix="7.2  Entire Agreement.",
     text="This Agreement constitutes the entire agreement between the Parties with respect to its "
          "subject matter and supersedes all prior discussions, representations, or agreements.")
body("", bold_prefix="7.3  Amendments.",
     text="This Agreement may not be modified except by a written amendment signed by both Parties.")
body("", bold_prefix="7.4  No Waiver.",
     text="Failure by either Party to enforce any provision of this Agreement shall not constitute "
          "a waiver of future enforcement of that provision.")
body("", bold_prefix="7.5  Severability.",
     text="If any provision of this Agreement is found unenforceable, the remaining provisions "
          "shall continue in full force and effect.")
body("", bold_prefix="7.6  Counterparts.",
     text="This Agreement may be executed in counterparts, including by electronic signature, "
          "each of which shall be deemed an original.")

# ============================================================
# SIGNATURES
# ============================================================
blank_line()
heading("SIGNATURES", level=1)
body(
    "By signing below, each Party agrees to be bound by the terms of this Agreement."
)
blank_line()

# Two-column signature block via tab stops
def sig_block(left_label, right_label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(6)
    tab_stop = OxmlElement('w:tab')
    tab_stop.set(qn('w:val'), 'left')
    tab_stop.set(qn('w:pos'), '4320')  # 3 inches
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tabs.append(tab_stop)
    pPr.append(tabs)
    r1 = p.add_run('_' * 38)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run('\t' + '_' * 38)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

def label_row(left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    tab_stop = OxmlElement('w:tab')
    tab_stop.set(qn('w:val'), 'left')
    tab_stop.set(qn('w:pos'), '4320')
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tabs.append(tab_stop)
    pPr.append(tabs)
    r1 = p.add_run(left)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2 = p.add_run('\t' + right)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

sig_block("Signature", "Signature")
label_row("Gleeson Consulting, LLC", "Kedron Cheerleading Association")

sig_block("Name", "Name")
label_row("Holly Gleeson", "Jenifer Friel")

sig_block("Title", "Title")
label_row("Principal", "Director")

sig_block("Date", "Date")
label_row("Date", "Date")

# ============================================================
# SAVE
# ============================================================
output = "/Users/hollygleeson/Dropbox/CheerCast/CheerCast_Pilot_Agreement.docx"
doc.save(output)
print(f"Saved: {output}")
